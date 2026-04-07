"""
md2docx.py — Markdown → Word (.docx) 转换脚本
依赖: python-docx  （pip install python-docx）

用法:
    python docs/md2docx.py <input.md> [output.docx]

若不指定输出路径，则在同目录下生成同名 .docx 文件。
"""

import sys
import re
from pathlib import Path


# ── 依赖检查 ────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("缺少依赖，请先安装: pip install python-docx")
    sys.exit(1)


# ── 辅助：为段落中的 run 添加代码字体 ────────────────────────────────────────
def _set_mono(run):
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


# ── 辅助：解析行内标记（bold / italic / inline-code / link） ─────────────────
INLINE_RE = re.compile(
    r"(`+)(.+?)\1"                          # inline code
    r"|(\*\*|__)(.+?)\3"                    # bold
    r"|(\*|_)(.+?)\5"                       # italic
    r"|\[([^\]]+)\]\(([^)]*)\)"            # link [text](url)
)

def _add_inline(para, text: str):
    """将含有 inline 标记的文本添加到段落，逐段拆分为 runs。"""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # 前面的普通文本
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        pos = m.end()

        if m.group(1):                      # inline code
            r = para.add_run(m.group(2))
            _set_mono(r)
        elif m.group(3):                    # bold
            r = para.add_run(m.group(4))
            r.bold = True
        elif m.group(5):                    # italic
            r = para.add_run(m.group(6))
            r.italic = True
        elif m.group(7) is not None:        # link — 只显示文字
            r = para.add_run(m.group(7))
            r.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
            r.underline = True

    if pos < len(text):
        para.add_run(text[pos:])


# ── 辅助：在文档末尾添加水平线 ──────────────────────────────────────────────
def _add_hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 辅助：解析 GFM 表格 ──────────────────────────────────────────────────────
def _parse_table(lines: list[str]):
    """返回 list[list[str]]，第 0 行是表头。分隔行自动跳过。"""
    rows = []
    for line in lines:
        if re.match(r"^\|?[\s\-:|]+\|", line):   # 分隔行
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline(para, cell_text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()  # 表后空行


# ── 核心转换 ─────────────────────────────────────────────────────────────────
def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()

    doc = Document()

    # 全局字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # ── 空行 ────────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── 水平线 ──────────────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            _add_hr(doc)
            i += 1
            continue

        # ── ATX 标题 (#  ##  ### …) ─────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).rstrip("#").strip()
            para = doc.add_heading("", level=min(level, 4))
            _add_inline(para, heading_text)
            i += 1
            continue

        # ── 围栏代码块 (``` … ```) ──────────────────────────────────────────
        if line.startswith("```") or line.startswith("~~~"):
            fence = line[:3]
            i += 1
            code_lines = []
            while i < n and not lines[i].startswith(fence):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束围栏
            code_text = "\n".join(code_lines)
            para = doc.add_paragraph()
            para.style = "No Spacing"
            run = para.add_run(code_text)
            _set_mono(run)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 给整段加浅灰背景
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F5F5")
            pPr.append(shd)
            doc.add_paragraph()
            continue

        # ── 表格 (GFM) ──────────────────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < n and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table(table_lines)
            _add_table(doc, rows)
            continue

        # ── 无序列表 (- / * / +) ────────────────────────────────────────────
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2  # 缩进层级
            style_name = "List Bullet" if indent == 0 else "List Bullet 2"
            para = doc.add_paragraph(style=style_name)
            _add_inline(para, m.group(2))
            i += 1
            continue

        # ── 有序列表 (1. 2. …) ──────────────────────────────────────────────
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            style_name = "List Number" if indent == 0 else "List Number 2"
            para = doc.add_paragraph(style=style_name)
            _add_inline(para, m.group(2))
            i += 1
            continue

        # ── 引用块 (>) ──────────────────────────────────────────────────────
        if line.startswith(">"):
            quote_text = re.sub(r"^>\s?", "", line)
            para = doc.add_paragraph(style="Intense Quote")
            _add_inline(para, quote_text)
            i += 1
            continue

        # ── 普通段落 ─────────────────────────────────────────────────────────
        # 合并连续非空行为同一段落
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if (not nxt.strip()
                    or nxt.startswith("#")
                    or nxt.startswith(">")
                    or nxt.startswith("```")
                    or nxt.startswith("~~~")
                    or re.match(r"^[-*+]\s", nxt)
                    or re.match(r"^\d+[.)]\s", nxt)
                    or ("|" in nxt and nxt.strip().startswith("|"))
                    or re.match(r"^[-*_]{3,}\s*$", nxt.strip())):
                break
            para_lines.append(nxt)
            i += 1

        para = doc.add_paragraph()
        _add_inline(para, " ".join(para_lines))

    doc.save(docx_path)
    print(f"已生成: {docx_path}")


# ── 入口 ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)

    md_file = Path(sys.argv[1])
    if not md_file.exists():
        print(f"文件不存在: {md_file}")
        sys.exit(1)

    docx_file = Path(sys.argv[2]) if len(sys.argv) >= 3 else md_file.with_suffix(".docx")
    convert(md_file, docx_file)